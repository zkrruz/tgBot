from __future__ import annotations

import re
import subprocess
import tempfile
from pathlib import Path

import pdfplumber
from docx import Document
from striprtf.striprtf import rtf_to_text


SUPPORTED_EXTENSIONS = {".pdf", ".doc", ".docx", ".rtf", ".txt"}


def normalize_text(text: str) -> str:
    text = text.replace("\x00", " ")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def extract_resume_text(path: Path) -> str:
    ext = path.suffix.lower()
    if ext not in SUPPORTED_EXTENSIONS:
        raise ValueError(f"Unsupported file format: {ext}")
    if ext == ".pdf":
        return _from_pdf(path)
    if ext == ".docx":
        return _from_docx(path)
    if ext == ".rtf":
        return _from_rtf(path)
    if ext == ".doc":
        return _from_doc(path)
    return normalize_text(path.read_text(encoding="utf-8", errors="ignore"))


def _from_pdf(path: Path) -> str:
    chunks: list[str] = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            chunks.append(page.extract_text() or "")
    return normalize_text("\n".join(chunks))


def _from_docx(path: Path) -> str:
    document = Document(str(path))
    parts = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text for cell in row.cells))
    return normalize_text("\n".join(parts))


def _from_rtf(path: Path) -> str:
    return normalize_text(rtf_to_text(path.read_text(encoding="utf-8", errors="ignore")))


def _from_doc(path: Path) -> str:
    for command in ("antiword", "catdoc"):
        try:
            result = subprocess.run(
                [command, str(path)],
                check=True,
                capture_output=True,
                text=True,
                timeout=20,
            )
            return normalize_text(result.stdout)
        except (FileNotFoundError, subprocess.SubprocessError):
            continue

    with tempfile.TemporaryDirectory() as tmp:
        out_dir = Path(tmp)
        try:
            subprocess.run(
                [
                    "soffice",
                    "--headless",
                    "--convert-to",
                    "txt:Text",
                    "--outdir",
                    str(out_dir),
                    str(path),
                ],
                check=True,
                capture_output=True,
                text=True,
                timeout=45,
            )
            txt_files = list(out_dir.glob("*.txt"))
            if txt_files:
                return normalize_text(txt_files[0].read_text(encoding="utf-8", errors="ignore"))
        except (FileNotFoundError, subprocess.SubprocessError):
            pass

    raise RuntimeError("Cannot read .doc file. Install antiword, catdoc, or LibreOffice on the server.")
